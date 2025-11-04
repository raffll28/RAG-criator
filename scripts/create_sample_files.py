"""
Script para criar arquivos de exemplo PDF e DOCX
"""

import fitz  # PyMuPDF
from docx import Document as DocxDoc
from docx.shared import Pt, Inches
from pathlib import Path


def create_pdf_report():
    """Cria um PDF de relatório mensal"""
    doc = fitz.open()
    page = doc.new_page()
    
    # Título
    page.insert_text((50, 50), "RELATÓRIO MENSAL DE VENDAS", fontsize=16)
    page.insert_text((50, 80), "Período: Outubro/2025", fontsize=12)
    page.insert_text((50, 95), "_" * 80, fontsize=10)
    
    # Conteúdo
    y = 120
    content = [
        "RESUMO EXECUTIVO",
        "",
        "Vendas Totais: R$ 1.245.000,00",
        "Crescimento: +15% em relação ao mês anterior",
        "Novos Clientes: 147",
        "Taxa de Conversão: 4,2%",
        "",
        "DESTAQUES DO MÊS",
        "",
        "• Lançamento de nova linha de produtos",
        "• Parceria estratégica com fornecedor internacional",
        "• Expansão para 3 novas regiões",
        "• Implementação de novo sistema de CRM",
        "",
        "ANÁLISE POR CATEGORIA",
        "",
        "Eletrônicos: R$ 450.000 (36%)",
        "Móveis: R$ 320.000 (26%)",
        "Vestuário: R$ 275.000 (22%)",
        "Outros: R$ 200.000 (16%)",
        "",
        "METAS PARA NOVEMBRO",
        "",
        "1. Aumentar vendas em 10%",
        "2. Reduzir tempo de entrega em 20%",
        "3. Melhorar NPS para 8,5",
        "4. Contratar 5 novos vendedores",
        "",
        "OBSERVAÇÕES",
        "",
        "O desempenho do mês foi excepcional, superando todas as expectativas.",
        "A equipe está motivada e preparada para os desafios do próximo período.",
    ]
    
    for line in content:
        page.insert_text((50, y), line, fontsize=10)
        y += 15
    
    # Salva o PDF
    output_path = Path("data/raw/relatorio_vendas.pdf")
    doc.save(output_path)
    doc.close()
    print(f"[OK] PDF criado: {output_path}")


def create_pdf_article():
    """Cria um PDF com artigo técnico"""
    doc = fitz.open()
    page = doc.new_page()
    
    # Título
    page.insert_text((50, 50), "Introdução ao Machine Learning", fontsize=16)
    page.insert_text((50, 75), "Por: Dr. Carlos Silva", fontsize=10)
    page.insert_text((50, 90), "_" * 80, fontsize=10)
    
    # Conteúdo
    y = 115
    content = [
        "O que é Machine Learning?",
        "",
        "Machine Learning (Aprendizado de Máquina) é uma área da Inteligência",
        "Artificial que permite que sistemas aprendam e melhorem automaticamente",
        "a partir da experiência, sem serem explicitamente programados.",
        "",
        "Tipos de Aprendizado",
        "",
        "1. Aprendizado Supervisionado",
        "   Utiliza dados rotulados para treinar modelos que fazem previsões.",
        "   Exemplos: Classificação de emails, previsão de preços.",
        "",
        "2. Aprendizado Não Supervisionado",
        "   Encontra padrões em dados não rotulados.",
        "   Exemplos: Segmentação de clientes, detecção de anomalias.",
        "",
        "3. Aprendizado por Reforço",
        "   O agente aprende através de tentativa e erro.",
        "   Exemplos: Jogos, robótica, sistemas de recomendação.",
        "",
        "Aplicações Práticas",
        "",
        "• Assistentes virtuais (Siri, Alexa)",
        "• Sistemas de recomendação (Netflix, Spotify)",
        "• Carros autônomos",
        "• Diagnóstico médico",
        "• Detecção de fraudes",
        "",
        "Ferramentas Populares",
        "",
        "• Python: Linguagem mais utilizada",
        "• TensorFlow: Framework do Google",
        "• PyTorch: Framework do Facebook",
        "• Scikit-learn: Biblioteca para ML tradicional",
        "",
        "Conclusão",
        "",
        "Machine Learning está transformando indústrias e criando novas",
        "oportunidades. O futuro promete avanços ainda mais impressionantes.",
    ]
    
    for line in content:
        if y > 750:  # Nova página se necessário
            page = doc.new_page()
            y = 50
        page.insert_text((50, y), line, fontsize=9)
        y += 13
    
    # Salva o PDF
    output_path = Path("data/raw/artigo_ml.pdf")
    doc.save(output_path)
    doc.close()
    print(f"[OK] PDF criado: {output_path}")


def create_docx_proposal():
    """Cria um DOCX de proposta comercial"""
    doc = DocxDoc()
    
    # Título
    title = doc.add_heading('PROPOSTA COMERCIAL', 0)
    title.alignment = 1  # Centralizado
    
    # Informações básicas
    doc.add_paragraph('Empresa: TechSolutions Ltda.')
    doc.add_paragraph('Data: 04 de Novembro de 2025')
    doc.add_paragraph('Validade: 30 dias')
    doc.add_paragraph('Versão: 1.0')
    
    # Seção 1
    doc.add_heading('1. APRESENTAÇÃO', 1)
    doc.add_paragraph(
        'A TechSolutions é especializada em desenvolvimento de soluções tecnológicas '
        'inovadoras para empresas de médio e grande porte. Com mais de 10 anos de '
        'experiência no mercado, já atendemos mais de 200 clientes em todo o Brasil.'
    )
    
    # Seção 2
    doc.add_heading('2. ESCOPO DO PROJETO', 1)
    doc.add_paragraph('O projeto consiste no desenvolvimento de um sistema de gestão empresarial (ERP) customizado, incluindo:')
    
    items = [
        'Módulo Financeiro (contas a pagar e receber)',
        'Módulo de Vendas (pedidos, notas fiscais)',
        'Módulo de Estoque (controle de produtos)',
        'Módulo de Relatórios (dashboards e analytics)',
        'Integração com sistemas existentes',
        'Treinamento da equipe'
    ]
    
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Seção 3
    doc.add_heading('3. CRONOGRAMA', 1)
    
    # Tabela de cronograma
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fase'
    hdr_cells[1].text = 'Atividade'
    hdr_cells[2].text = 'Prazo'
    
    data = [
        ('1', 'Levantamento de Requisitos', '2 semanas'),
        ('2', 'Desenvolvimento', '8 semanas'),
        ('3', 'Testes', '2 semanas'),
        ('4', 'Implantação e Treinamento', '1 semana'),
    ]
    
    for i, (fase, atividade, prazo) in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = fase
        cells[1].text = atividade
        cells[2].text = prazo
    
    # Seção 4
    doc.add_heading('4. INVESTIMENTO', 1)
    doc.add_paragraph('Desenvolvimento do Sistema: R$ 85.000,00')
    doc.add_paragraph('Treinamento: R$ 5.000,00')
    doc.add_paragraph('Suporte (12 meses): R$ 12.000,00')
    doc.add_paragraph('')
    
    p = doc.add_paragraph('TOTAL: ')
    p.add_run('R$ 102.000,00').bold = True
    
    # Seção 5
    doc.add_heading('5. CONDIÇÕES DE PAGAMENTO', 1)
    doc.add_paragraph('30% na assinatura do contrato')
    doc.add_paragraph('40% na entrega do sistema')
    doc.add_paragraph('30% após homologação final')
    
    # Seção 6
    doc.add_heading('6. CONTATO', 1)
    doc.add_paragraph('Email: contato@techsolutions.com.br')
    doc.add_paragraph('Telefone: (11) 3456-7890')
    doc.add_paragraph('Website: www.techsolutions.com.br')
    
    # Salva o DOCX
    output_path = Path("data/raw/proposta_comercial.docx")
    doc.save(output_path)
    print(f"[OK] DOCX criado: {output_path}")


def create_docx_manual():
    """Cria um DOCX de manual do usuário"""
    doc = DocxDoc()
    
    # Título
    title = doc.add_heading('MANUAL DO USUÁRIO', 0)
    title.alignment = 1
    
    subtitle = doc.add_heading('Sistema de Gestão Empresarial', 2)
    subtitle.alignment = 1
    
    doc.add_paragraph('')
    
    # Capítulo 1
    doc.add_heading('1. INTRODUÇÃO', 1)
    doc.add_paragraph(
        'Bem-vindo ao Sistema de Gestão Empresarial. Este manual foi criado para '
        'ajudá-lo a utilizar todas as funcionalidades do sistema de forma eficiente.'
    )
    
    # Capítulo 2
    doc.add_heading('2. PRIMEIROS PASSOS', 1)
    
    doc.add_heading('2.1 Acesso ao Sistema', 2)
    doc.add_paragraph('1. Abra seu navegador web')
    doc.add_paragraph('2. Digite o endereço: https://sistema.empresa.com')
    doc.add_paragraph('3. Insira seu usuário e senha')
    doc.add_paragraph('4. Clique em "Entrar"')
    
    doc.add_heading('2.2 Interface Principal', 2)
    doc.add_paragraph(
        'Após o login, você verá o dashboard principal com os seguintes elementos:'
    )
    doc.add_paragraph('• Menu lateral com todas as funcionalidades', style='List Bullet')
    doc.add_paragraph('• Barra superior com perfil do usuário', style='List Bullet')
    doc.add_paragraph('• Área central com indicadores e gráficos', style='List Bullet')
    
    # Capítulo 3
    doc.add_heading('3. FUNCIONALIDADES PRINCIPAIS', 1)
    
    doc.add_heading('3.1 Cadastro de Clientes', 2)
    doc.add_paragraph('Para cadastrar um novo cliente:')
    doc.add_paragraph('1. Acesse o menu "Clientes"')
    doc.add_paragraph('2. Clique em "Novo Cliente"')
    doc.add_paragraph('3. Preencha os dados obrigatórios')
    doc.add_paragraph('4. Clique em "Salvar"')
    
    doc.add_heading('3.2 Emissão de Notas Fiscais', 2)
    doc.add_paragraph('Para emitir uma nota fiscal:')
    doc.add_paragraph('1. Acesse o menu "Vendas" > "Notas Fiscais"')
    doc.add_paragraph('2. Clique em "Nova NF"')
    doc.add_paragraph('3. Selecione o cliente')
    doc.add_paragraph('4. Adicione os produtos')
    doc.add_paragraph('5. Clique em "Emitir"')
    
    # Tabela de atalhos
    doc.add_heading('4. ATALHOS DE TECLADO', 1)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light List Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Atalho'
    hdr_cells[1].text = 'Função'
    
    shortcuts = [
        ('Ctrl + N', 'Novo registro'),
        ('Ctrl + S', 'Salvar'),
        ('Ctrl + P', 'Imprimir'),
        ('F5', 'Atualizar'),
        ('Esc', 'Cancelar'),
    ]
    
    for i, (atalho, funcao) in enumerate(shortcuts, 1):
        cells = table.rows[i].cells
        cells[0].text = atalho
        cells[1].text = funcao
    
    # Seção de suporte
    doc.add_paragraph('')
    doc.add_heading('5. SUPORTE TÉCNICO', 1)
    doc.add_paragraph('Em caso de dúvidas ou problemas:')
    doc.add_paragraph('Email: suporte@sistema.com')
    doc.add_paragraph('Telefone: 0800 123 4567')
    doc.add_paragraph('Chat online: Disponível das 8h às 18h')
    
    # Salva o DOCX
    output_path = Path("data/raw/manual_usuario.docx")
    doc.save(output_path)
    print(f"[OK] DOCX criado: {output_path}")


if __name__ == "__main__":
    print("Criando arquivos de exemplo...\n")
    
    create_pdf_report()
    create_pdf_article()
    create_docx_proposal()
    create_docx_manual()
    
    print("\n[OK] Todos os arquivos foram criados com sucesso!")

