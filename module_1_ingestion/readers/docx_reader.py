"""
DOCX Reader - Leitor de arquivos Microsoft Word

Suporta: .docx
Usa python-docx para extração de texto e metadados
"""

from pathlib import Path
from typing import Union
import hashlib
from docx import Document as DocxDocument
from loguru import logger

from .base_reader import BaseReader, Document


class DOCXReader(BaseReader):
    """
    Reader para arquivos Microsoft Word (.docx).
    
    Extrai:
    - Texto de parágrafos
    - Texto de tabelas
    - Metadados do documento (autor, título, etc.)
    """
    
    def __init__(self, include_tables: bool = True):
        """
        Inicializa o DOCXReader.
        
        Args:
            include_tables: Se True, extrai texto de tabelas também
        """
        super().__init__()
        self.include_tables = include_tables
        self.supported_extensions = [".docx"]
    
    def read(self, file_path: Union[Path, str]) -> Document:
        """
        Lê um arquivo DOCX e retorna um Document.
        
        Args:
            file_path: Caminho do arquivo DOCX (Path ou string)
            
        Returns:
            Document contendo o conteúdo e metadados
            
        Raises:
            FileNotFoundError: Se o arquivo não existir
            ValueError: Se o arquivo não puder ser lido
        """
        # Converte para Path se necessário
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        logger.debug(f"Reading DOCX file: {file_path}")
        
        # Valida o arquivo
        self._validate_file_exists(file_path)
        
        # Abre o documento
        try:
            docx_document = DocxDocument(file_path)
        except Exception as e:
            raise ValueError(f"Failed to open DOCX {file_path}: {e}")
        
        # Extrai texto dos parágrafos
        content_parts = []
        
        for paragraph in docx_document.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)
        
        # Extrai texto das tabelas se habilitado
        table_count = 0
        if self.include_tables:
            for table in docx_document.tables:
                table_count += 1
                table_text = self._extract_table_text(table)
                if table_text:
                    content_parts.append(f"\n[Tabela {table_count}]\n{table_text}")
        
        content = "\n".join(content_parts)
        
        # Metadados básicos do arquivo
        metadata = self._extract_basic_metadata(file_path)
        
        # Metadados específicos do DOCX
        core_properties = docx_document.core_properties
        
        metadata["docx_title"] = core_properties.title or ""
        metadata["docx_author"] = core_properties.author or ""
        metadata["docx_subject"] = core_properties.subject or ""
        metadata["docx_keywords"] = core_properties.keywords or ""
        metadata["docx_comments"] = core_properties.comments or ""
        metadata["docx_last_modified_by"] = core_properties.last_modified_by or ""
        metadata["docx_revision"] = core_properties.revision or 0
        
        # Estatísticas do documento
        metadata["paragraphs_count"] = len(docx_document.paragraphs)
        metadata["tables_count"] = table_count
        metadata["sections_count"] = len(docx_document.sections)
        
        # Metadados de conteúdo
        metadata["content_length"] = len(content)
        metadata["char_count"] = len(content)
        metadata["lines_count"] = len(content.splitlines()) if content else 0
        metadata["word_count"] = len(content.split()) if content else 0
        metadata["content_hash"] = hashlib.md5(content.encode()).hexdigest()
        
        # Determina se permite vazio
        allow_empty = len(content) == 0
        
        if allow_empty:
            logger.warning(f"DOCX {file_path.name} has no extractable text.")
        
        # Cria o documento
        document = Document(
            content=content,
            metadata=metadata,
            source=str(file_path.absolute()),
            allow_empty=allow_empty
        )
        
        logger.info(
            f"Successfully read DOCX {file_path.name} "
            f"({metadata['paragraphs_count']} paragraphs, "
            f"{metadata['tables_count']} tables, "
            f"{metadata['word_count']} words)"
        )
        
        return document
    
    def _extract_table_text(self, table) -> str:
        """
        Extrai texto de uma tabela do Word.
        
        Args:
            table: Objeto Table do python-docx
            
        Returns:
            Texto formatado da tabela
        """
        rows_text = []
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells_text)
            if row_text.strip():
                rows_text.append(row_text)
        
        return "\n".join(rows_text)

